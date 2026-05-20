import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "_pylibs"))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

OUT_DIR = os.path.dirname(__file__)

BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_ACCENT = RGBColor(0x00, 0x5A, 0x9E)
BRAND_LIGHT_BG = RGBColor(0xF2, 0xF4, 0xF7)
BRAND_GRAY = RGBColor(0x6B, 0x70, 0x80)
BRAND_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_checkbox_xml(paragraph):
    """Add a Word-native checkbox content control (structured document tag)."""
    run = paragraph.add_run()
    fldChar1 = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin">'
        f'<w:ffData><w:name w:val="Check1"/><w:enabled/>'
        f'<w:calcOnExit w:val="0"/><w:checkBox><w:sizeAuto/>'
        f'<w:default w:val="0"/></w:checkBox></w:ffData></w:fldChar>'
    )
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> FORMCHECKBOX </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def add_text_field(paragraph, width_chars=60):
    """Add a Word form text field (legacy form field) that is fillable."""
    run = paragraph.add_run()
    fldChar1 = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin">'
        f'<w:ffData><w:name w:val="Text1"/><w:enabled/>'
        f'<w:calcOnExit w:val="0"/><w:textInput>'
        f'<w:maxLength w:val="{width_chars * 3}"/>'
        f'</w:textInput></w:ffData></w:fldChar>'
    )
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> FORMTEXT </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar_sep)

    run4 = paragraph.add_run("\u200B")
    run4.font.color.rgb = BRAND_GRAY
    run4.font.size = Pt(11)

    run5 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar2)


def setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BRAND_DARK

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BRAND_ACCENT
    p.space_after = Pt(4)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(11)
        run2.font.color.rgb = BRAND_GRAY
        run2.italic = True
        p2.space_after = Pt(16)


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_ACCENT


def add_instruction(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_GRAY
    run.italic = True
    p.space_after = Pt(8)


def add_label_and_field(doc, label):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(label + "  ")
    run.bold = True
    run.font.size = Pt(11)
    add_text_field(p)


def add_multiline_field(doc, label, lines=4):
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    for _ in range(lines):
        p2 = doc.add_paragraph()
        p2.space_after = Pt(2)
        add_text_field(p2, width_chars=100)


def add_checkbox_item(doc, text):
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    add_checkbox_xml(p)
    run = p.add_run("  " + text)
    run.font.size = Pt(11)


def add_table(doc, headers, rows_data=None, num_blank_rows=0):
    """Create a formatted table. rows_data is list of lists. num_blank_rows adds empty fillable rows."""
    total_rows = 1 + (len(rows_data) if rows_data else 0) + num_blank_rows
    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_WHITE
        set_cell_shading(cell, "005A9E")

    row_idx = 1
    if rows_data:
        for row in rows_data:
            for i, val in enumerate(row):
                cell = table.rows[row_idx].cells[i]
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None
            row_idx += 1

    for r in range(row_idx, total_rows):
        for i in range(len(headers)):
            cell = table.rows[r].cells[i]
            cell.text = ""
            add_text_field(cell.paragraphs[0], width_chars=40)
        if r % 2 == 0:
            for i in range(len(headers)):
                set_cell_shading(table.rows[r].cells[i], "F2F4F7")

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# WORKSHEET 1: PLANNING
# ──────────────────────────────────────────────────────────────
def build_ws1():
    doc = setup_doc()
    add_title(doc, "WORKSHEET 1: PLANNING",
              "Effective AI Coding  |  Module 1")

    add_instruction(doc,
        "Complete this worksheet before writing any prompts. "
        "Every answer here becomes part of your seed file.")

    # 1.1
    add_section_header(doc, "1.1  Project Overview")
    add_instruction(doc,
        "Write a single paragraph (3\u20135 sentences). "
        "What is it? Who is it for? What problem does it solve?")
    add_multiline_field(doc, "Your answer:", lines=5)

    # 1.2
    add_section_header(doc, "1.2  Architecture Decisions")
    add_instruction(doc,
        "Fill in each row with your project\u2019s choices. "
        "Do not leave any blank. If unsure, research and pick one.")
    for label in ["Cloud Provider", "Compute", "Operating System",
                  "Instance / Server Type", "Storage", "Database",
                  "Network / Access", "Port"]:
        add_label_and_field(doc, label)

    # 1.3
    add_section_header(doc, "1.3  Tech Stack")
    add_instruction(doc,
        'Be specific. "Python" is not enough \u2014 write "Python 3.12 with FastAPI."')
    for label in ["Backend", "Frontend", "Database",
                  "External APIs / Services", "Authentication Method"]:
        add_label_and_field(doc, label)

    add_page_break(doc)

    # 1.4
    add_section_header(doc, "1.4  Feature Scope \u2014 What Is In")
    add_instruction(doc,
        "List every feature that MUST be in this iteration. Be specific about behavior.\n"
        'Bad: "User management"\n'
        'Good: "Admins can create, deactivate, and delete user accounts. '
        'Standard users cannot access admin functions."')
    add_table(doc, ["#", "Feature Description"], num_blank_rows=10)

    # 1.5
    add_section_header(doc, "1.5  Feature Scope \u2014 What Is Out")
    add_instruction(doc,
        "List features you are deliberately deferring. "
        "This prevents the AI from adding things you did not ask for.")
    add_table(doc, ["#", "Deferred Feature", "Reason for Deferral"], num_blank_rows=5)

    # 1.6
    add_section_header(doc, "1.6  Roles and Access Rules")
    add_instruction(doc,
        "Who uses your application? What can each role do?")
    add_table(doc, ["Role Name", "Permissions / Can Do", "Restricted From"],
              num_blank_rows=4)
    add_label_and_field(doc, "What requires authentication?")
    add_label_and_field(doc, "What is accessible without login?")

    add_page_break(doc)

    # 1.7
    add_section_header(doc, "1.7  UX / Design Rules")
    add_instruction(doc,
        "Write 5\u20138 plain-language rules about how the application should look "
        "and behave. These become constraints for the AI.")
    add_table(doc, ["#", "Design Rule"], num_blank_rows=8)

    path = os.path.join(OUT_DIR, "Worksheet 1 - Planning.docx")
    doc.save(path)
    print(f"  Created: {path}")


# ──────────────────────────────────────────────────────────────
# WORKSHEET 2: GOAL PROTOTYPING
# ──────────────────────────────────────────────────────────────
def build_ws2():
    doc = setup_doc()
    add_title(doc, "WORKSHEET 2: GOAL PROTOTYPING",
              "Effective AI Coding  |  Module 2")

    add_instruction(doc,
        "Complete this worksheet after creating your seed file and "
        "generating the interactive wireframe prototype.")

    # 2.1
    add_section_header(doc, "2.1  Seed File Creation Checklist")
    add_instruction(doc, "Check each item as you add it to your seed file.")
    for item in [
        "Project overview paragraph",
        "Architecture and infrastructure table",
        "Tech stack table",
        "Authentication and user management rules",
        "Feature descriptions with specific behavior",
        "Data model / what gets stored",
        "UI layout description (pages, navigation, key elements)",
        "Deployment constraints (port, network, environment)",
        "What is deferred / not included",
    ]:
        add_checkbox_item(doc, item)
    add_label_and_field(doc, "Total sections in your seed file")

    # 2.2
    add_section_header(doc, "2.2  Wireframe Prompt")
    add_instruction(doc,
        "Write the exact prompt you will give the AI to generate the wireframe. "
        "Include a reference to the seed file and specify fake data.")
    add_multiline_field(doc, "Your prompt:", lines=4)

    add_page_break(doc)

    # 2.3
    add_section_header(doc, "2.3  Prototype Review Checklist")
    add_instruction(doc,
        "After the AI generates the wireframe, walk through it and check each item.")
    for item in [
        "Login flow works (enter credentials, see the app)",
        "Every page from the seed file has a corresponding screen",
        "Navigation between pages works",
        "All user roles are represented (can you see the difference?)",
        "Data displays look realistic",
        "Action buttons are present (create, edit, delete, export, etc.)",
        "Error states are represented (what happens when something fails?)",
        "The overall look and feel matches your design rules",
    ]:
        add_checkbox_item(doc, item)

    add_section_header(doc, "Issues Found")
    add_table(doc, ["#", "Issue Description", "Severity (High/Med/Low)"],
              num_blank_rows=6)

    # 2.4
    add_section_header(doc, "2.4  Stakeholder Feedback")
    add_instruction(doc,
        "Show the prototype to at least one other person. Record their feedback.")
    add_label_and_field(doc, "Reviewer name / role")
    add_table(doc, ["#", "Feedback Item", "Action Needed?"], num_blank_rows=5)

    path = os.path.join(OUT_DIR, "Worksheet 2 - Goal Prototyping.docx")
    doc.save(path)
    print(f"  Created: {path}")


# ──────────────────────────────────────────────────────────────
# WORKSHEET 3: FEATURE CHANGES
# ──────────────────────────────────────────────────────────────
def build_ws3():
    doc = setup_doc()
    add_title(doc, "WORKSHEET 3: FEATURE CHANGES",
              "Effective AI Coding  |  Module 3")

    add_instruction(doc,
        "Complete this worksheet after the prototype review and "
        "before instructing the AI to build the application.")

    # 3.1
    add_section_header(doc, "3.1  Change Log")
    add_instruction(doc,
        "List every change from the prototype review and stakeholder feedback. "
        "Type: NEW (new feature), CHANGE (modify existing), or REMOVE.")
    add_table(doc, ["#", "Type", "Description", "Seed File Section"],
              num_blank_rows=8)

    # 3.2
    add_section_header(doc, "3.2  Design Questions for the AI")
    add_instruction(doc,
        "List questions you need the AI to help answer before building. "
        "Record the options it presents and YOUR decision.")
    for i in range(1, 4):
        p = doc.add_paragraph()
        p.space_before = Pt(10)
        run = p.add_run(f"Question {i}")
        run.bold = True
        run.font.size = Pt(11)
        add_label_and_field(doc, "Your question")
        add_label_and_field(doc, "AI option a)")
        add_label_and_field(doc, "AI option b)")
        add_label_and_field(doc, "AI option c)")
        add_label_and_field(doc, "Your decision")
        add_label_and_field(doc, "Rationale")

    add_page_break(doc)

    # 3.3
    add_section_header(doc, "3.3  Seed File Update Verification")
    add_instruction(doc,
        "After updating the seed file, re-read it completely and check each item.")
    for item in [
        "Every change from 3.1 is reflected in the seed file",
        "No contradictions exist between old and new sections",
        "Every feature has enough detail for the AI to build it",
        "The seed file reads coherently from start to finish",
    ]:
        add_checkbox_item(doc, item)
    add_label_and_field(doc, "Total section count")

    path = os.path.join(OUT_DIR, "Worksheet 3 - Feature Changes.docx")
    doc.save(path)
    print(f"  Created: {path}")


# ──────────────────────────────────────────────────────────────
# WORKSHEET 4: PROJECT CREATION
# ──────────────────────────────────────────────────────────────
def build_ws4():
    doc = setup_doc()
    add_title(doc, "WORKSHEET 4: PROJECT CREATION",
              "Effective AI Coding  |  Module 4")

    add_instruction(doc,
        "Complete this worksheet during and after the AI generates the application code.")

    # 4.1
    add_section_header(doc, "4.1  Build Prompt")
    add_instruction(doc,
        "Write the exact prompt you will give the AI to build the application.")
    add_multiline_field(doc, "Your prompt:", lines=4)

    # 4.2
    add_section_header(doc, "4.2  Generated File Review")
    add_instruction(doc,
        "After the AI generates the project, list every file it created. "
        "Mark each as reviewed after you have read and understood it.")
    add_table(doc, ["File Path", "Reviewed?", "Notes"], num_blank_rows=12)

    add_page_break(doc)

    # 4.3
    add_section_header(doc, "4.3  Initial Smoke Test")
    add_instruction(doc,
        "Run through these basic checks immediately after code generation.")
    for item in [
        "The project starts without syntax errors",
        "The main page loads in a browser",
        "Login / authentication works",
        "Basic navigation between pages works",
        "At least one core feature works end-to-end",
    ]:
        add_checkbox_item(doc, item)
    add_section_header(doc, "Issues Found")
    add_table(doc, ["#", "Issue Description"], num_blank_rows=5)

    # 4.4
    add_section_header(doc, "4.4  Bug Report Log")
    add_instruction(doc,
        "Track each bug you find and how you reported it to the AI.")
    add_table(doc, ["#", "Symptom (what you saw)", "Fix Applied?", "Verified?"],
              num_blank_rows=8)

    path = os.path.join(OUT_DIR, "Worksheet 4 - Project Creation.docx")
    doc.save(path)
    print(f"  Created: {path}")


# ──────────────────────────────────────────────────────────────
# WORKSHEET 5: DEPLOYMENT AND SETUP
# ──────────────────────────────────────────────────────────────
def build_ws5():
    doc = setup_doc()
    add_title(doc, "WORKSHEET 5: DEPLOYMENT AND SETUP",
              "Effective AI Coding  |  Module 5")

    add_instruction(doc,
        "Complete this worksheet while deploying the application to the target environment.")

    # 5.1
    add_section_header(doc, "5.1  Environment Preparation Checklist")
    for item in [
        "Server / instance provisioned",
        "Operating system is the expected version",
        "Network / firewall rules configured",
        "System dependencies installed",
        "Application user and directories created",
        "Application code uploaded",
        "Runtime dependencies installed (pip, npm, etc.)",
    ]:
        add_checkbox_item(doc, item)
    add_multiline_field(doc, "Notes on issues encountered:", lines=3)

    # 5.2
    add_section_header(doc, "5.2  Secrets and Configuration")
    add_instruction(doc,
        "List every secret or credential your application needs.")
    add_table(doc, ["Secret Name", "Stored Where", "Generated By"], num_blank_rows=5)
    add_checkbox_item(doc, "All secret files have restricted permissions (600 or 640)")
    add_checkbox_item(doc, "No secrets are in source code or version control")

    add_page_break(doc)

    # 5.3
    add_section_header(doc, "5.3  Startup Verification")
    for item in [
        "Service starts without errors",
        "Application responds on the expected port",
        "Logs are being written (stdout and/or file)",
        "Database was created / initialized",
        "Default admin account exists and can log in",
    ]:
        add_checkbox_item(doc, item)

    # 5.4
    add_section_header(doc, "5.4  Initial Configuration")
    for item in [
        "Default admin password changed",
        "External service credentials configured (API keys, etc.)",
        "Application settings configured (thresholds, timeouts, policies)",
        "Additional user accounts created",
        "All configured users can log in",
    ]:
        add_checkbox_item(doc, item)

    path = os.path.join(OUT_DIR, "Worksheet 5 - Deployment and Setup.docx")
    doc.save(path)
    print(f"  Created: {path}")


# ──────────────────────────────────────────────────────────────
# WORKSHEET 6: TESTING AND VALIDATION
# ──────────────────────────────────────────────────────────────
def build_ws6():
    doc = setup_doc()
    add_title(doc, "WORKSHEET 6: TESTING AND VALIDATION",
              "Effective AI Coding  |  Module 6")

    add_instruction(doc,
        "Complete this worksheet after the application is deployed and running.")

    # 6.1
    add_section_header(doc, "6.1  Test Plan")
    add_instruction(doc,
        "Convert each feature from your seed file into a test case. "
        'Format: "When [action], expect [result]."')
    add_table(doc, ["#", "When [action]...", "Expect [result]..."], num_blank_rows=12)

    # 6.2
    add_section_header(doc, "6.2  Functional Test Results")
    add_table(doc, ["Test #", "Pass / Fail", "Notes"], num_blank_rows=12)

    add_page_break(doc)

    # 6.3
    add_section_header(doc, "6.3  Edge Case Tests")
    for item in [
        "Bad input / empty fields are handled gracefully",
        "Unauthorized access attempts are blocked",
        "External service unavailable is handled with a clear message",
        "Resource limits (disk, memory) behave as designed",
        "Session timeout works correctly",
        "Concurrent usage does not cause errors",
    ]:
        add_checkbox_item(doc, item)

    # 6.4
    add_section_header(doc, "6.4  Security Checklist")
    for item in [
        "Credentials are encrypted at rest",
        "No sensitive data in application logs",
        "Session cookies are httponly / secure",
        "Role-based access is enforced on the backend (not just the UI)",
        "File permissions on secrets are restrictive",
        "Default credentials have been changed",
    ]:
        add_checkbox_item(doc, item)

    path = os.path.join(OUT_DIR, "Worksheet 6 - Testing and Validation.docx")
    doc.save(path)
    print(f"  Created: {path}")


# ──────────────────────────────────────────────────────────────
# WORKSHEET 7: HANDOFF AND ITERATION
# ──────────────────────────────────────────────────────────────
def build_ws7():
    doc = setup_doc()
    add_title(doc, "WORKSHEET 7: HANDOFF AND ITERATION",
              "Effective AI Coding  |  Module 7")

    add_instruction(doc,
        "Complete this worksheet as you prepare the project for team use and plan the next iteration.")

    # 7.1
    add_section_header(doc, "7.1  Documentation Package Checklist")
    for item in [
        "Project requirements / seed file",
        "Setup and deployment guide",
        "Project steps overview",
        "Architecture decisions document",
        "Troubleshooting guide",
        "All documents are up to date with the final state of the code",
    ]:
        add_checkbox_item(doc, item)
    add_label_and_field(doc, "Location of documentation")

    # 7.2
    add_section_header(doc, "7.2  Knowledge Transfer Checklist")
    for item in [
        "Demonstrated the application to the team",
        "Walked through admin functions",
        "Explained backup and recovery procedures",
        "Explained how to update the application",
        "Explained how to rotate secrets / credentials",
        "Answered team questions",
    ]:
        add_checkbox_item(doc, item)
    add_section_header(doc, "Team Members Trained")
    add_table(doc, ["#", "Name", "Role"], num_blank_rows=5)

    add_page_break(doc)

    # 7.3
    add_section_header(doc, "7.3  Iteration 2 Backlog")
    add_instruction(doc,
        "Collect feedback and prioritize features for the next iteration.")
    add_table(doc, ["Priority", "Feature / Improvement", "Source"],
              num_blank_rows=8)

    # 7.4
    add_section_header(doc, "7.4  Retrospective")
    add_instruction(doc,
        "Answer honestly \u2014 this improves your next project.")
    add_multiline_field(doc, "What worked well in this process?", lines=3)
    add_multiline_field(doc, "What was frustrating or inefficient?", lines=3)
    add_multiline_field(doc, "What would you do differently next time?", lines=3)
    add_multiline_field(doc, "How much time did the planning phase save during the build phase?", lines=2)

    path = os.path.join(OUT_DIR, "Worksheet 7 - Handoff and Iteration.docx")
    doc.save(path)
    print(f"  Created: {path}")


if __name__ == "__main__":
    print("Building worksheets...")
    build_ws1()
    build_ws2()
    build_ws3()
    build_ws4()
    build_ws5()
    build_ws6()
    build_ws7()
    print("Done. 7 worksheets created.")
